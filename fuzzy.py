Penyebab utama dari masalah pada gambar yang Anda unggah adalah karena instruksi *floating position* (memaksa tabel ke paling bawah margin) berbenturan dengan sistem *auto-repeat header* pada halaman 2. Hal ini membuat tata letak Microsoft Word menjadi *error* dan memotong baris tabel tanda tangan menjadi dua bagian yang berjauhan.

**Solusi Paling Aman dan Anti-Gagal (Bulletproof):**
Kita akan membuang instruksi *floating* yang bermasalah tersebut, dan mengubah arsitektur tabel tanda tangan menjadi **Tabel 1 Baris (1-Row Table)**.

Judul ("CHECKED / EXAMINED") dan isian ("Name & Date") akan digabungkan ke dalam **satu sel yang sama**. Dengan menerapkan kode `prevent_row_split` pada 1 baris ini, Microsoft Word **secara mutlak tidak akan bisa memisahkan** judul dan isian tanda tangan ke halaman yang berbeda. Jika tidak muat di halaman 1, maka *seluruh blok* tanda tangan akan turun bersama-sama secara utuh dan rapi ke halaman 2.

Berikut adalah *source code* lengkap yang sudah diperbaiki secara menyeluruh:

### Full Source Code Python (Streamlit + `python-docx`)

```python
import numpy as np
import streamlit as st
import pandas as pd
import math
from datetime import datetime
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# Konfigurasi Halaman Utama
st.set_page_config(page_title="Engineering Tools & LPI Report", layout="wide")

# =======================================================================================
# FUNGSI PEMBANTU FORMATTING XML (UNTUK PAGINATION & LAYOUT TINGKAT LANJUT)
# =======================================================================================
def set_cell_margins(cell, top=100, bottom=100, start=100, end=100):
    """Mengatur padding di dalam cell tabel (dalam dxa)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', start), ('w:right', end)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, fill_hex):
    """Mengatur warna latar belakang cell"""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def prevent_row_split(row):
    """Mencegah baris tabel terpotong di antara dua halaman"""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(r'<w:cantSplit {}/>'.format(nsdecls('w'))))

def set_repeat_header(row):
    """Mengatur baris tabel agar otomatis diulang di halaman berikutnya"""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(r'<w:tblHeader {}/>'.format(nsdecls('w'))))

def add_xml_field(run, field_name):
    """Menyisipkan field kode dinamis MS Word (PAGE / NUMPAGES)"""
    fldChar1 = parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w')))
    instrText = parse_xml(r'<w:instrText {} xml:space="preserve"> {} </w:instrText>'.format(nsdecls('w'), field_name))
    fldChar2 = parse_xml(r'<w:fldChar {} w:fldCharType="separate"/>'.format(nsdecls('w')))
    fldChar3 = parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w')))
    run._r.extend([fldChar1, instrText, fldChar2, fldChar3])

# =======================================================================================
# FUNGSI UTAMA GENERATE .DOCX
# =======================================================================================
def generate_docx_report(logo_left_bytes, logo_right_top_bytes, logo_right_bottom_bytes, 
                         client, project, equipment, auto_form_no, date_str, doc_no, rev_no,
                         drawing_no, standard, description, penetrant_method, removal_method, 
                         brand_name, penetrant_type, developer_type, cleaner_type, 
                         surface_prep, time_exam, scope_exam, data_df, logo_width_inch, logo_height_inch,
                         dict_joint_photos):
    
    doc = Document()
    
    # Atur Margin Halaman Cetak Standar Internasional A4
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(2.4)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.header_distance = Inches(0.4)
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)

    # -----------------------------------------------------------------------------------
    # 1. INTEGRASI KOP SURAT KE NATIVE WORD HEADER
    # -----------------------------------------------------------------------------------
    header = doc.sections[0].header
    for p in header.paragraphs:
        p.text = ""
        
    kop_table = header.add_table(3, 4, Inches(6.77))
    kop_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    kop_table.style = 'Table Grid'
    
    col_widths = [Inches(1.85), Inches(3.07), Inches(0.85), Inches(1.00)]
    for row in kop_table.rows:
        prevent_row_split(row)
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    # Merging Kolom Kop
    cell_left = kop_table.cell(0, 0).merge(kop_table.cell(1, 0))          
    cell_center = kop_table.cell(0, 1).merge(kop_table.cell(1, 1))        
    cell_right_top = kop_table.cell(0, 2).merge(kop_table.cell(0, 3))     
    cell_right_bottom = kop_table.cell(1, 2).merge(kop_table.cell(1, 3))  

    # Logo Kiri
    cell_left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_left = cell_left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo_left_bytes is not None:
        p_left.add_run().add_picture(logo_left_bytes, width=Inches(logo_width_inch), height=Inches(logo_height_inch))
    else:
        p_left.add_run("[ LOGO KIRI ]").font.color.rgb = RGBColor(160, 160, 160)

    # Judul Tengah
    cell_center.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_center = cell_center.paragraphs[0]
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_center = p_center.add_run(str(project).upper())
    run_center.bold = True
    run_center.font.size = Pt(11)
    run_center.font.name = 'Arial'

    # Logo Kanan Atas
    cell_right_top.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_rt = cell_right_top.paragraphs[0]
    p_rt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo_right_top_bytes is not None:
        p_rt.add_run().add_picture(logo_right_top_bytes, width=Inches(logo_width_inch), height=Inches(logo_height_inch))
    else:
        p_rt.add_run("[ LOGO KANAN ATAS ]").font.color.rgb = RGBColor(160, 160, 160)

    # Logo Kanan Bawah
    cell_right_bottom.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_rb = cell_right_bottom.paragraphs[0]
    p_rb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo_right_bottom_bytes is not None:
        p_rb.add_run().add_picture(logo_right_bottom_bytes, width=Inches(logo_width_inch), height=Inches(logo_height_inch))
    else:
        p_rb.add_run("[ LOGO KANAN BAWAH ]").font.color.rgb = RGBColor(160, 160, 160)

    # Baris Metadata Kop
    meta_cells = kop_table.rows[2].cells
    for cell in meta_cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, top=60, bottom=60, start=100, end=100)

    meta_cells[0].paragraphs[0].text = f"Date: {date_str}"
    meta_cells[1].paragraphs[0].text = f"Doc No.: {doc_no if doc_no else '-'}"
    meta_cells[2].paragraphs[0].text = f"Rev. {rev_no}"
    
    p_page = meta_cells[3].paragraphs[0]
    p_page.text = "Page "
    add_xml_field(p_page.add_run(), "PAGE")
    p_page.add_run(" of ")
    add_xml_field(p_page.add_run(), "NUMPAGES")
    
    for cell in meta_cells:
        p = cell.paragraphs[0]
        p.runs[0].font.size = Pt(9.5)
        p.runs[0].font.name = 'Arial'
    meta_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # -----------------------------------------------------------------------------------
    # BODY CONTENT AREA
    # -----------------------------------------------------------------------------------
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("FINAL LIQUID PENETRANT INSPECTION REPORT")
    h1_run.bold = True
    h1_run.font.size = Pt(13)
    h1_run.font.name = 'Arial'
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_after = Pt(12)

    # Tabel Informasi Utama
    table_header = doc.add_table(rows=5, cols=4)
    table_header.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_header.style = 'Table Grid'
    info_widths = [Inches(1.5), Inches(2.0), Inches(1.5), Inches(1.77)]
    
    headers_data = [
        ["CLIENT", client, "FORM NO", auto_form_no],
        ["PROJECT", project, "DATE", date_str],
        ["EQUIPMENT / SYSTEM", equipment, "PENETRANT METHOD", f"☑ {penetrant_method}"],
        ["DRAWING NO", drawing_no, "REMOVAL METHOD", f"☑ {removal_method}"],
        ["STANDARD / DESC", f"{standard} / {description}", "BRAND & MATERIALS", f"{brand_name} ({penetrant_type}/{developer_type})"]
    ]
    
    for row_idx, row_data in enumerate(headers_data):
        row_cells = table_header.rows[row_idx].cells
        prevent_row_split(table_header.rows[row_idx])
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].width = info_widths[col_idx]
            row_cells[col_idx].text = str(text)
            row_cells[col_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, start=100, end=100)
            p = row_cells[col_idx].paragraphs[0]
            p.runs[0].font.size = Pt(9.5)
            p.runs[0].font.name = 'Arial'
            if col_idx in [0, 2]:
                p.runs[0].font.bold = True
                set_cell_background(row_cells[col_idx], "F8F9FA")

    # Parameter Pengujian
    p_param = doc.add_paragraph()
    p_param.paragraph_format.space_before = Pt(10)
    p_param.paragraph_format.space_after = Pt(12)
    params = [("Surface Prep: ", f"☑ {surface_prep}   |   "), ("Time of Exam: ", f"☑ {time_exam}   |   "), ("Scope: ", f"☑ {scope_exam}")]
    for lbl, val in params:
        p_param.add_run(lbl).bold = True
        p_param.runs[-1].font.size = Pt(9.5)
        p_param.runs[-1].font.name = 'Arial'
        p_param.add_run(val).font.size = Pt(9.5)
        p_param.runs[-1].font.name = 'Arial'

    # Tabel Hasil Utama
    h2 = doc.add_paragraph()
    h2.add_run("Inspection Results Table").bold = True
    h2.runs[-1].font.size = Pt(11)
    h2.runs[-1].font.name = 'Arial'
    
    table_res = doc.add_table(rows=1, cols=7)
    table_res.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_res.style = 'Table Grid'
    res_widths = [Inches(1.8), Inches(0.7), Inches(1.1), Inches(0.5), Inches(0.6), Inches(1.2), Inches(0.87)]
    
    hdr_cells = table_res.rows[0].cells
    set_repeat_header(table_res.rows[0])
    prevent_row_split(table_res.rows[0])
    headers_col = ["PART NAME", "WELD NO", "THICKNESS (MM)", "ACC", "REJECT", "TYPES OF DISCONTINUITIES", "REMARKS"]
    for i, title_text in enumerate(headers_col):
        hdr_cells[i].width = res_widths[i]
        hdr_cells[i].text = title_text
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(hdr_cells[i], top=100, bottom=100, start=60, end=60)
        set_cell_background(hdr_cells[i], "EFEFEF")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.name = 'Arial'

    for index, row in data_df.iterrows():
        new_row = table_res.add_row()
        prevent_row_split(new_row)
        row_cells = new_row.cells
        for idx in range(7):
            row_cells[idx].width = res_widths[idx]
            row_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(row_cells[idx], top=80, bottom=80, start=60, end=60)
            
        row_cells[0].text = str(row["PART NAME"])
        row_cells[1].text = str(row["WELD NO"])
        row_cells[2].text = f"{row['THICKNESS (MM)']:.2f}" if isinstance(row['THICKNESS (MM)'], (int, float)) else str(row['THICKNESS (MM)'])
        row_cells[3].text = "☑" if row["RESULT"] == "ACC" else "☐"
        row_cells[4].text = "☑" if row["RESULT"] == "REJECT" else "☐"
        row_cells[5].text = str(row["TYPES OF DISCONTINUITIES"])
        row_cells[6].text = str(row["REMARKS"])
        
        for col_idx in [1, 2, 3, 4]:
            row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for idx in range(7):
            p = row_cells[idx].paragraphs[0]
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.name = 'Arial'
            
    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    # -----------------------------------------------------------------------------------
    # PERBAIKAN TOTAL: TABEL TANDA TANGAN ANTI-SPLIT (1 BARIS SAJA)
    # Menyatukan judul dan isian dalam 1 sel agar tidak akan pernah terpisah oleh page break
    # -----------------------------------------------------------------------------------
    table_sig = doc.add_table(rows=1, cols=4)
    table_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_sig.style = None 
    prevent_row_split(table_sig.rows[0]) # Mengunci tabel 1 baris ini secara absolut
    
    sig_widths = [Inches(1.69), Inches(1.69), Inches(1.69), Inches(1.70)]
    sig_titles = ["CHECKED / EXAMINED", "REVIEWED", "REVIEWED / WITNESSED", "REVIEWED / WITNESSED"]
    
    for idx, title_text in enumerate(sig_titles):
        cell = table_sig.rows[0].cells[idx]
        cell.width = sig_widths[idx]
        set_cell_margins(cell, top=60, bottom=60, start=60, end=60)
        
        # Paragraf 1: Judul di Center
        p_t = cell.paragraphs[0]
        p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER 
        p_t.paragraph_format.keep_with_next = True # Menggandeng teks bawahnya
        r_t = p_t.add_run(title_text)
        r_t.bold = True
        r_t.font.size = Pt(9)
        r_t.font.name = 'Arial'
        
        # Paragraf 2: Isian Tanda Tangan, Nama, Tanggal di Kiri
        p_b = cell.add_paragraph()
        p_b.alignment = WD_ALIGN_PARAGRAPH.LEFT 
        r_b = p_b.add_run("\n\n\n\n_______________________\nName:\nDate:")
        r_b.font.size = Pt(9)
        r_b.font.name = 'Arial'

    # -----------------------------------------------------------------------------------
    # AREA LAMPIRAN FOTO: OTOMATIS LOCK ASPECT RATIO 100% (NATIVE SIZE)
    # -----------------------------------------------------------------------------------
    has_any_photo = any(
        (weld_id in dict_joint_photos) and 
        (dict_joint_photos[weld_id]["red"] is not None or dict_joint_photos[weld_id]["dev"] is not None)
        for weld_id in data_df["WELD NO"].astype(str).tolist()
    )
    
    if has_any_photo:
        doc.add_page_break() 
        
        h_app = doc.add_paragraph()
        h_app.add_run("APPENDIX: PHOTOGRAPHIC DOCUMENTATION PER-JOINT").bold = True
        h_app.runs[-1].font.size = Pt(12)
        h_app.runs[-1].font.name = 'Arial'
        h_app.paragraph_format.space_after = Pt(14)
        
        for _, row in data_df.iterrows():
            weld_no_str = str(row["WELD NO"])
            part_name_str = str(row["PART NAME"])
            
            if weld_no_str not in dict_joint_photos:
                continue
                
            photo_data = dict_joint_photos[weld_no_str]
            if photo_data["red"] is None and photo_data["dev"] is None:
                continue
                
            p_jtitle = doc.add_paragraph()
            p_jtitle.paragraph_format.space_before = Pt(8)
            p_jtitle.paragraph_format.space_after = Pt(4)
            p_jtitle.paragraph_format.keep_with_next = True 
            
            run_jt = p_jtitle.add_run(f"■ Joint No: {weld_no_str} ({part_name_str})")
            run_jt.bold = True
            run_jt.font.size = Pt(10)
            run_jt.font.name = 'Arial'
            
            photo_table = doc.add_table(rows=2, cols=2)
            photo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            photo_table.style = 'Table Grid'
            
            prevent_row_split(photo_table.rows[0])
            prevent_row_split(photo_table.rows[1])
            for cell in photo_table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = True
            
            photo_table.rows[0].cells[0].width = Inches(3.38)
            photo_table.rows[0].cells[1].width = Inches(3.39)
            photo_table.rows[1].cells[0].width = Inches(3.38)
            photo_table.rows[1].cells[1].width = Inches(3.39)
            
            # Slot Foto Kiri: Red Apply 
            cell_r_img = photo_table.cell(0, 0)
            cell_r_img.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell_r_img, top=180, bottom=180, start=180, end=180)
            
            if photo_data["red"] is not None:
                cell_r_img.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell_r_img.paragraphs[0].add_run().add_picture(photo_data["red"])
                cell_r_img.paragraphs[0].paragraph_format.keep_with_next = True
            else:
                cell_r_img.paragraphs[0].text = "[ Foto Red Apply Tidak Tersedia ]"
                cell_r_img.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell_r_img.paragraphs[0].paragraph_format.keep_with_next = True
                
            cell_r_cap = photo_table.cell(1, 0)
            set_cell_background(cell_r_cap, "F8F9FA")
            cell_r_cap.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_cap_run = cell_r_cap.paragraphs[0].add_run(f"Joint {weld_no_str}: Penetrant Application (Red Apply)")
            r_cap_run.font.size = Pt(8.5)
            r_cap_run.font.name = 'Arial'
            
            # Slot Foto Kanan: Developer Apply
            cell_d_img = photo_table.cell(0, 1)
            cell_d_img.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell_d_img, top=180, bottom=180, start=180, end=180)
            
            if photo_data["dev"] is not None:
                cell_d_img.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell_d_img.paragraphs[0].add_run().add_picture(photo_data["dev"])
                cell_d_img.paragraphs[0].paragraph_format.keep_with_next = True
            else:
                cell_d_img.paragraphs[0].text = "[ Foto Developer Tidak Tersedia ]"
                cell_d_img.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell_d_img.paragraphs[0].paragraph_format.keep_with_next = True
                
            cell_d_cap = photo_table.cell(1, 1)
            set_cell_background(cell_d_cap, "F8F9FA")
            cell_d_cap.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            d_cap_run = cell_d_cap.paragraphs[0].add_run(f"Joint {weld_no_str}: Developer Application")
            d_cap_run.font.size = Pt(8.5)
            d_cap_run.font.name = 'Arial'
            
            doc.add_paragraph().paragraph_format.space_after = Pt(6)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# =======================================================================================
# MENU UTAMA APLIKASI STREAMLIT
# =======================================================================================
st.sidebar.title("Navigation Menu")
main_menu = st.sidebar.radio("PILIH MODUL APLIKASI:", ["LIQUID PENETRANT REPORT", "CONSUMABLE CALCULATE"])

st.sidebar.markdown("---")
st.sidebar.info("Aplikasi Integrasi: LPI Generator & Volume Calculator.")

if main_menu == "LIQUID PENETRANT REPORT":
    st.title("📋 Liquid Penetrant Inspection Report Generator")
    st.write("Aplikasi untuk men-generate report hasil pengujian Liquid Penetrant dengan lampiran foto asli proporsional per-joint.")

    # SIDEBAR KONTROL SLIDER KOP LOGO
    st.sidebar.subheader("📐 Ukuran Komponen Master (.docx)")
    logo_w_setting = st.sidebar.slider("Lebar Semua Logo Kop (Inchi)", min_value=0.8, max_value=2.0, value=1.3, step=0.05)
    logo_h_setting = st.sidebar.slider("Tinggi Semua Logo Kop (Inchi)", min_value=0.8, max_value=2.0, value=1.3, step=0.05)
    st.sidebar.info("💡 Catatan: Foto lampiran (attachment) sekarang di-lock otomatis 100% sesuai ukuran & aspect ratio file aslinya.")

    # 1. BLOK UNGGAH LOGO PERUSAHAAN (KOP ATAS)
    st.subheader("🖼️ Konfigurasi Kop Surat (Multi-Logo Dinamis)")
    log_col1, log_col2, log_col3 = st.columns(3)
    with log_col1:
        up_logo_left = st.file_uploader("1. Logo Kiri (Main Contractor / Owner)", type=["png", "jpg", "jpeg"])
    with log_col2:
        up_logo_rt = st.file_uploader("2. Logo Kanan Atas (Project Brand)", type=["png", "jpg", "jpeg"])
    with log_col3:
        up_logo_rb = st.file_uploader("3. Logo Kanan Bawah (Sub-Consultant / QA)", type=["png", "jpg", "jpeg"])

    st.markdown("---")

    current_date = datetime.now()
    date_str = current_date.strftime("12-06-2026") 
    date_display_str = current_date.strftime("%d %B %Y").upper() 
    month_year_slug = current_date.strftime("%B/%Y").upper() 
    auto_form_no = f"05/LPI/DAMAC/{month_year_slug}" 

    st.subheader("📝 Header & Document Control Information")
    col1, col2, col3 = st.columns(3)
    with col1:
        client = st.text_input("Client", value="PT. TRAKINDO UTAMA")
        project = st.text_input("Project", value="DAMAC DIGITAL JKT01 DAY 2")
        equipment = st.text_input("Equipment / System", value="FUEL PIPE") 
    with col2:
        st.text_input("Form No (Otomatis)", value=auto_form_no, disabled=True)
        st.text_input("Date (Otomatis)", value=date_str, disabled=True)
        drawing_no = st.text_input("Drawing No", value="ISO-JKT01-003") 
    with col3:
        doc_no_input = st.text_input("Doc No. (Kop)", value="DOC-DAMAC-LPI-002")
        rev_no_input = st.text_input("Rev. No (Kop)", value="0")
        standard = st.text_input("Standard", value="ASME B31.3")
        description = st.text_input("Description", value="TYPE 1")

    st.markdown("---")
    st.subheader("🔍 Inspection Parameters")
    col_p1, col_p2, col_p3 = st.columns(3)
    with col_p1:
        penetrant_method = st.radio("Penetrant Method", ["VISIBLE", "FLUORECENT"], index=0)
        removal_method = st.radio("Removal Method", ["SOLVENT REMOVABLE", "WATER WASHABLE", "POST EMULSIFIEBLE"], index=0)
    with col_p2:
        brand_name = st.text_input("Brand Name", value="MAGNAFLUX")
        penetrant_type = st.text_input("Penetrant Code", value="SPL-SP2")
        developer_type = st.text_input("Developer Code", value="SKD-S2")
        cleaner_type = st.text_input("Cleaner Code", value="SKC-S")
    with col_p3:
        surface_prep = st.radio("Surface Preparation", ["AS WELDED", "MACHINING", "GRINDING", "OTHER"], index=0)
        time_exam = st.radio("Time of Examination", ["AFTER WELDING", "AFTER HYDROTEST", "AFTER PWHT", "OTHER"], index=0)
        scope_exam = st.radio("Scope of Examination", ["BASE METAL", "WELD METAL", "BACK CHIPPING", "OTHER"], index=1)

    st.markdown("---")
    st.subheader("📊 Weld Inspection Results Data")

    part_options_1 = ["PIPE", "PLATE", "ELBOW", "TEE", "FLANGE"]
    part_options_2 = ["FLANGE", "ELBOW", "EQUAL TEE", "PIPE", "VALVE"]
    discontinuity_options = ["Crack", "Porosity", "Slag Inclusion", "Incomplete Fusion", "Incomplete Penetration", "Undercut", "Linear Indication", "Rounded Indication"]

    if 'weld_data' not in st.session_state:
        st.session_state.weld_data = pd.DataFrame([
            {"PART NAME": "PIPE – FLANGE", "WELD NO": "1", "THICKNESS (MM)": 3.91, "RESULT": "ACC", "TYPES OF DISCONTINUITIES": "-", "REMARKS": "-"},
            {"PART NAME": "PIPE – ELBOW", "WELD NO": "2", "THICKNESS (MM)": 3.91, "RESULT": "ACC", "TYPES OF DISCONTINUITIES": "-", "REMARKS": "-"},
            {"PART NAME": "PIPE – EQUAL TEE", "WELD NO": "6", "THICKNESS (MM)": 3.91, "RESULT": "ACC", "TYPES OF DISCONTINUITIES": "-", "REMARKS": "-"}
        ])

    with st.expander("➕ Tambah Baris Hasil Las Baru"):
        c1_a, c1_b, c2, c3, c4 = st.columns([2, 2, 1.5, 1.5, 2])
        with c1_a: part_1 = st.selectbox("Pilihan 1 (Base)", part_options_1)
        with c1_b: part_2 = st.selectbox("Pilihan 2 (Joint)", part_options_2)
        with c2: new_weld_no = st.text_input("Weld No", value="3")
        with c3: new_thickness = st.number_input("Thickness (mm)", value=3.91, step=0.01)
        with c4: new_result = st.selectbox("Result", ["ACC", "REJECT"])
        
        new_discontinuity = "-"
        new_remarks = "Repair Required" if new_result == "REJECT" else "-"
        if new_result == "REJECT":
            new_discontinuity = st.selectbox("Types of Discontinuities", discontinuity_options)

        if st.button("Masukkan ke Tabel"):
            combined_part_name = f"{part_1} – {part_2}"
            new_row = {"PART NAME": combined_part_name, "WELD NO": new_weld_no, "THICKNESS (MM)": new_thickness,
                       "RESULT": new_result, "TYPES OF DISCONTINUITIES": new_discontinuity, "REMARKS": new_remarks}
            st.session_state.weld_data = pd.concat([st.session_state.weld_data, pd.DataFrame([new_row])], ignore_index=True)
            st.success(f"Data baru '{combined_part_name}' berhasil ditambahkan ke tabel!")

    all_combined_options = [f"{p1} – {p2}" for p1 in part_options_1 for p2 in part_options_2]
    existing_parts = st.session_state.weld_data["PART NAME"].unique().tolist()
    final_part_config_options = list(set(all_combined_options + existing_parts))

    edited_weld_df = st.data_editor(
        st.session_state.weld_data, 
        num_rows="dynamic", 
        use_container_width=True,
        column_config={
            "PART NAME": st.column_config.SelectboxColumn("PART NAME", options=final_part_config_options, required=True),
            "RESULT": st.column_config.SelectboxColumn("RESULT", options=["ACC", "REJECT"], required=True),
            "TYPES OF DISCONTINUITIES": st.column_config.SelectboxColumn("TYPES OF DISCONTINUITIES", options=["-"] + discontinuity_options, required=True)
        }
    )

    # 2. INTERFACE UPLOAD FOTO PARALEL PER-JOINT SECARA DINAMIS
    st.markdown("---")
    st.subheader("📸 Upload Dokumentasi Foto Lapangan Per-Joint")
    st.write("Silakan buka expander di bawah ini untuk memasukkan foto spesifik pada masing-masing nomor joint:")
    
    master_joint_photos = {}
    
    for idx, row in edited_weld_df.iterrows():
        w_id = str(row["WELD NO"])
        p_name = str(row["PART NAME"])
        
        with st.expander(f"Joint No: {w_id} ({p_name})"):
            img_col1, img_col2 = st.columns(2)
            with img_col1:
                f_red = st.file_uploader(f"Foto Red Apply (Joint {w_id})", type=["png","jpg","jpeg"], key=f"red_{w_id}_{idx}")
                f_red_bytes = io.BytesIO(f_red.getvalue()) if f_red else None
            with img_col2:
                f_dev = st.file_uploader(f"Foto Dev Apply (Joint {w_id})", type=["png","jpg","jpeg"], key=f"dev_{w_id}_{idx}")
                f_dev_bytes = io.BytesIO(f_dev.getvalue()) if f_dev else None
                
            master_joint_photos[w_id] = {"red": f_red_bytes, "dev": f_dev_bytes}

    st.markdown("---")
    
    # Pengolahan Byte Kop Logo
    logo_l_io = io.BytesIO(up_logo_left.getvalue()) if up_logo_left else None
    logo_rt_io = io.BytesIO(up_logo_rt.getvalue()) if up_logo_rt else None
    logo_rb_io = io.BytesIO(up_logo_rb.getvalue()) if up_logo_rb else None

    col_btn1, col_btn2 = st.columns([2, 8])
    with col_btn1:
        generate_layout = st.button("🚀 Generate Screen Layout")
        
    with col_btn2:
        docx_buffer = generate_docx_report(
            logo_left_bytes=logo_l_io, logo_right_top_bytes=logo_rt_io, logo_right_bottom_bytes=logo_rb_io,
            client=client, project=project, equipment=equipment, auto_form_no=auto_form_no, 
            date_str=date_str, doc_no=doc_no_input, rev_no=rev_no_input,
            drawing_no=drawing_no, standard=standard, description=description, 
            penetrant_method=penetrant_method, removal_method=removal_method,
            brand_name=brand_name, penetrant_type=penetrant_type, developer_type=developer_type,
            cleaner_type=cleaner_type, surface_prep=surface_prep, time_exam=time_exam,
            scope_exam=scope_exam, data_df=edited_weld_df, 
            logo_width_inch=logo_w_setting, logo_height_inch=logo_h_setting,
            dict_joint_photos=master_joint_photos
        )
        
        st.download_button(
            label="📥 Download Official Word Report (.docx)",
            data=docx_buffer,
            file_name=f"LPI_Report_{auto_form_no.replace('/', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    if generate_layout:
        st.markdown("## 📄 PREVIEW REPORT LAYOUT (PAGE HEADER MODEL)")
        st.markdown(
            f"""
            <div style="border:2px solid #333; padding:15px; border-radius:5px; background-color:#FAFAFA">
                <table style="width:100%; border-collapse:collapse; border:none;">
                    <tr>
                        <td style="width:27%; text-align:center; border-right:1px solid #ccc; padding:10px;"><b>[ LOGO KIRI ]</b></td>
                        <td style="width:46%; text-align:center; border-right:1px solid #ccc; padding:10px;"><h4>{project.upper()}</h4></td>
                        <td style="width:27%; text-align:center; padding:10px;"><b>[ LOGO KANAN ATAS ]</b><hr style="margin:5px 0;"><b>[ LOGO KANAN BAWAH ]</b></td>
                    </tr>
                </table>
            </div>
            """, 
            unsafe_allow_html=True
        )
        
        st.markdown("<br>", unsafe_allow_html=True)
        header_summary = pd.DataFrame({
            "FIELD A": ["CLIENT", "PROJECT", "EQUIPMENT / SYSTEM", "DRAWING NO", "STANDARD / DESC"],
            "VALUE A": [client, project, equipment, drawing_no, f"{standard} / {description}"],
            "FIELD B": ["FORM NO", "DATE", "PENETRANT METHOD", "REMOVAL METHOD", "BRAND & MATERIALS"],
            "VALUE B": [auto_form_no, date_display_str, f"☑ {penetrant_method}", f"☑ {removal_method}", f"{brand_name} ({penetrant_type}/{developer_type})"]
        })
        st.table(header_summary)
        
        st.write("#### Inspection Results Table")
        display_df = edited_weld_df.copy()
        display_df["ACC"] = display_df["RESULT"].apply(lambda x: "☑" if x == "ACC" else "☐")
        display_df["REJECT"] = display_df["RESULT"].apply(lambda x: "☑" if x == "REJECT" else "☐")
        display_df = display_df[["PART NAME", "WELD NO", "THICKNESS (MM)", "ACC", "REJECT", "TYPES OF DISCONTINUITIES", "REMARKS"]]
        st.dataframe(display_df, use_container_width=True)
        
        st.write("#### Signatures")
        sig_cols = st.columns(4)
        sig_titles_preview = ["CHECKED / EXAMINED", "REVIEWED", "REVIEWED / WITNESSED", "REVIEWED / WITNESSED"]
        for i, title in enumerate(sig_titles_preview):
            with sig_cols[i]:
                st.write(f"**<center>{title}</center>**", unsafe_allow_html=True)
                st.write("\n\n\n__________________")
                st.write("Name:")
                st.write("Date:")
                
        # Preview Lampiran Foto Per-Joint di Web
        st.write(f"#### 📸 Attachment Preview: Per-Joint Photos (Locked Aspect Ratio 100%)")
        for key_w, media in master_joint_photos.items():
            if media["red"] or media["dev"]:
                st.write(f"**Joint Sambungan Las No: {key_w}**")
                p_col1, p_col2 = st.columns(2)
                with p_col1:
                    if media["red"]: st.image(media["red"], width=180, caption=f"Red Apply Joint {key_w}")
                with p_col2:
                    if media["dev"]: st.image(media["dev"], width=180, caption=f"Dev Apply Joint {key_w}")

```
